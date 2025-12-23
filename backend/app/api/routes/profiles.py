"""
Profile Management API Routes
"""

import os
from typing import List

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.database import get_db
from app import utils
from app.models.profile import Profile
from app.models.job import JobApplication, JobStatus
from app.schemas.profile import (
    ProfileCreate,
    ProfileUpdate,
    ProfileResponse,
    ProfileInternalResponse,
    ProfileListResponse,
    ProfileWithStats,
    ProfileStats,
)
from app.api.helpers import get_profile_or_404, validate_work_experience_index

router = APIRouter()


@router.get("", response_model=ProfileListResponse)
async def list_profiles(
    skip: int = 0,
    limit: int = 100,
    active_only: bool = True,
    db: AsyncSession = Depends(get_db),
):
    """List all profiles with optional filtering."""
    query = select(Profile)
    if active_only:
        query = query.where(Profile.is_active == True)
    query = query.offset(skip).limit(limit)

    result = await db.execute(query)
    profiles = result.scalars().all()

    # Get total count
    count_query = select(func.count(Profile.id))
    if active_only:
        count_query = count_query.where(Profile.is_active == True)
    total = await db.scalar(count_query) or 0

    return ProfileListResponse(
        profiles=[ProfileResponse.model_validate(p) for p in profiles],
        total=total,
    )


@router.get("/internal/all", response_model=list[ProfileInternalResponse])
async def get_all_profiles_internal(
    active_only: bool = True,
    db: AsyncSession = Depends(get_db),
):
    query = select(Profile)
    if active_only:
        query = query.where(Profile.is_active == True)
    
    result = await db.execute(query)
    profiles = result.scalars().all()
    
    return [ProfileInternalResponse.model_validate(p) for p in profiles]


@router.post("", response_model=ProfileResponse, status_code=status.HTTP_201_CREATED)
async def create_profile(
    profile_data: ProfileCreate,
    db: AsyncSession = Depends(get_db),
):
    """Create a new profile."""
    # Check for duplicate email
    existing = await db.scalar(
        select(Profile).where(Profile.email == profile_data.email)
    )
    if existing:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"Profile with email {profile_data.email} already exists",
        )

    # Compute display name
    display_name = f"{profile_data.first_name} {profile_data.last_name}"

    # Create profile with all new fields
    profile = Profile(
        # Name fields
        first_name=profile_data.first_name,
        middle_name=profile_data.middle_name,
        last_name=profile_data.last_name,
        preferred_first_name=profile_data.preferred_first_name,
        name=display_name,
        
        # Contact
        email=profile_data.email,
        phone=profile_data.phone,
        location=profile_data.location,
        preferred_password=profile_data.preferred_password,
        
        # Detailed Address for Job Applications
        address_1=profile_data.address_1,
        address_2=profile_data.address_2,
        county=profile_data.county,
        city=profile_data.city,
        state=profile_data.state,
        country=profile_data.country,
        zip_code=profile_data.zip_code,
        
        # Online Presence
        linkedin_url=profile_data.linkedin_url,
        github_url=profile_data.github_url,
        portfolio_url=profile_data.portfolio_url,
        
        # Demographics & Work Preferences
        gender=profile_data.gender,
        nationality=profile_data.nationality,
        veteran_status=profile_data.veteran_status,
        disability_status=profile_data.disability_status,
        willing_to_travel=profile_data.willing_to_travel,
        willing_to_relocate=profile_data.willing_to_relocate,
        primary_language=profile_data.primary_language,
        
        # Experience
        work_experience=[exp.model_dump() for exp in profile_data.work_experience],
        education=[edu.model_dump() for edu in profile_data.education],
        skills=profile_data.skills,
        custom_fields=profile_data.custom_fields,
        cover_letter_template=profile_data.cover_letter_template,
    )

    db.add(profile)
    await db.flush()
    await db.refresh(profile)

    return ProfileResponse.model_validate(profile)


@router.get("/{profile_id}", response_model=ProfileWithStats)
async def get_profile(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)
    stats = await _get_profile_stats(db, profile_id)
    response = ProfileWithStats.model_validate(profile)
    response.stats = stats
    return response


@router.put("/{profile_id}", response_model=ProfileResponse)
async def update_profile(
    profile_id: str,
    profile_data: ProfileUpdate,
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)

    # Update fields
    update_data = profile_data.model_dump(exclude_unset=True)

    # Handle nested objects
    if "work_experience" in update_data and update_data["work_experience"]:
        update_data["work_experience"] = [
            exp.model_dump() if hasattr(exp, "model_dump") else exp
            for exp in update_data["work_experience"]
        ]
    if "education" in update_data and update_data["education"]:
        update_data["education"] = [
            edu.model_dump() if hasattr(edu, "model_dump") else edu
            for edu in update_data["education"]
        ]

    for field, value in update_data.items():
        setattr(profile, field, value)

    # Update computed name if first/last name changed
    if "first_name" in update_data or "last_name" in update_data:
        profile.name = f"{profile.first_name} {profile.last_name}"

    await db.flush()
    await db.refresh(profile)

    return ProfileResponse.model_validate(profile)


@router.delete("/{profile_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_profile(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)
    await db.delete(profile)


@router.post("/{profile_id}/resume", response_model=ProfileResponse)
async def upload_resume(
    profile_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)

    # Validate file type
    allowed_types = ["application/pdf", "application/msword",
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and Word documents are allowed",
        )

    # Check file size
    max_size = settings.max_resume_size_mb * 1024 * 1024
    content = await file.read()
    if len(content) > max_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File size exceeds {settings.max_resume_size_mb}MB limit",
        )

    # Save file with original filename in profile-specific directory
    import re
    safe_name = re.sub(r'[^\w\s-]', '', profile.name).strip().replace(' ', '_')
    if not safe_name:
        safe_name = profile_id
    
    resume_dir = os.path.join(settings.storage_path, "resumes", safe_name)
    os.makedirs(resume_dir, exist_ok=True)

    original_filename = file.filename or "resume.pdf"
    file_path = os.path.join(resume_dir, original_filename)

    with open(file_path, "wb") as f:
        f.write(content)

    # Update profile
    profile.resume_path = file_path
    await db.flush()
    await db.refresh(profile)

    return ProfileResponse.model_validate(profile)


@router.delete("/{profile_id}/resume", response_model=ProfileResponse)
async def delete_resume(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)

    if not profile.resume_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No resume to delete",
        )

    if os.path.exists(profile.resume_path):
        os.remove(profile.resume_path)

    profile.resume_path = None
    await db.flush()
    await db.refresh(profile)

    return ProfileResponse.model_validate(profile)


@router.post("/{profile_id}/cover-letter-template", response_model=ProfileResponse)
async def upload_cover_letter_template(
    profile_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    import re
    
    profile = await get_profile_or_404(db, profile_id)

    allowed_types = [
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only DOC and DOCX files are allowed",
        )

    content = await file.read()

    safe_name = re.sub(r'[^\w\s-]', '', profile.name).strip().replace(' ', '_')
    if not safe_name:
        safe_name = profile_id
    
    template_dir = os.path.join(settings.storage_path, "cover_letter_templates", safe_name)
    os.makedirs(template_dir, exist_ok=True)

    original_filename = file.filename or "cover_letter_template.docx"
    file_path = os.path.join(template_dir, original_filename)

    with open(file_path, "wb") as f:
        f.write(content)

    profile.cover_letter_template_path = file_path
    await db.flush()
    await db.refresh(profile)

    return ProfileResponse.model_validate(profile)


@router.delete("/{profile_id}/cover-letter-template", response_model=ProfileResponse)
async def delete_cover_letter_template(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)

    if not profile.cover_letter_template_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No cover letter template to delete",
        )

    if os.path.exists(profile.cover_letter_template_path):
        os.remove(profile.cover_letter_template_path)

    profile.cover_letter_template_path = None
    await db.flush()
    await db.refresh(profile)

    return ProfileResponse.model_validate(profile)


@router.get("/{profile_id}/cover-letter-template/file")
async def get_cover_letter_template_file(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    from fastapi.responses import FileResponse
    
    profile = await get_profile_or_404(db, profile_id)
    
    if not profile.cover_letter_template_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No cover letter template uploaded",
        )
    
    if not os.path.exists(profile.cover_letter_template_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Cover letter template file not found",
        )
    
    filename = os.path.basename(profile.cover_letter_template_path)
    
    return FileResponse(
        path=profile.cover_letter_template_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"inline; filename=\"{filename}\""},
    )


@router.post("/{profile_id}/cover-letter-template/generate")
async def generate_cover_letter(
    profile_id: str,
    content: str,
    db: AsyncSession = Depends(get_db),
):
    from docx import Document
    import shutil
    import re
    
    profile = await get_profile_or_404(db, profile_id)
    
    if not profile.cover_letter_template_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No cover letter template uploaded",
        )
    
    if not os.path.exists(profile.cover_letter_template_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Cover letter template file not found",
        )
    
    safe_name = re.sub(r'[^\w\s-]', '', profile.name).strip().replace(' ', '_')
    if not safe_name:
        safe_name = profile_id
    
    generated_dir = os.path.join(settings.storage_path, "generated_cover_letters", safe_name)
    os.makedirs(generated_dir, exist_ok=True)
    
    generation_id = utils.generate_uuid()[:8]
    working_docx = os.path.join(generated_dir, f"cover_letter_{generation_id}.docx")
    output_pdf = os.path.join(generated_dir, f"cover_letter_{generation_id}.pdf")
    
    shutil.copy2(profile.cover_letter_template_path, working_docx)
    
    doc = Document(working_docx)
    
    for paragraph in doc.paragraphs:
        if "{{content}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{content}}", content)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{content}}" in cell.text:
                    cell.text = cell.text.replace("{{content}}", content)
    
    doc.save(working_docx)
    
    result_path = working_docx
    result_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    try:
        from docx2pdf import convert
        convert(working_docx, output_pdf)
        
        if os.path.exists(output_pdf):
            os.remove(working_docx)
            result_path = output_pdf
            result_type = "application/pdf"
    except Exception:
        pass
    
    return {
        "generation_id": generation_id,
        "file_path": result_path,
        "file_type": result_type,
    }


@router.get("/{profile_id}/cover-letter-template/generated/{generation_id}")
async def get_generated_cover_letter(
    profile_id: str,
    generation_id: str,
    db: AsyncSession = Depends(get_db),
):
    from fastapi.responses import FileResponse
    import re
    
    profile = await get_profile_or_404(db, profile_id)
    
    safe_name = re.sub(r'[^\w\s-]', '', profile.name).strip().replace(' ', '_')
    if not safe_name:
        safe_name = profile_id
    
    generated_dir = os.path.join(settings.storage_path, "generated_cover_letters", safe_name)
    
    pdf_path = os.path.join(generated_dir, f"cover_letter_{generation_id}.pdf")
    docx_path = os.path.join(generated_dir, f"cover_letter_{generation_id}.docx")
    
    if os.path.exists(pdf_path):
        return FileResponse(
            path=pdf_path,
            media_type="application/pdf",
            headers={"Content-Disposition": "inline; filename=\"cover_letter.pdf\""},
        )
    elif os.path.exists(docx_path):
        return FileResponse(
            path=docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "inline; filename=\"cover_letter.docx\""},
        )
    
    raise HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail="Generated cover letter not found",
    )


@router.delete("/{profile_id}/cover-letter-template/generated/{generation_id}")
async def delete_generated_cover_letter(
    profile_id: str,
    generation_id: str,
    db: AsyncSession = Depends(get_db),
):
    import re
    
    profile = await get_profile_or_404(db, profile_id)
    
    safe_name = re.sub(r'[^\w\s-]', '', profile.name).strip().replace(' ', '_')
    if not safe_name:
        safe_name = profile_id
    
    generated_dir = os.path.join(settings.storage_path, "generated_cover_letters", safe_name)
    
    pdf_path = os.path.join(generated_dir, f"cover_letter_{generation_id}.pdf")
    docx_path = os.path.join(generated_dir, f"cover_letter_{generation_id}.docx")
    
    deleted = False
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
        deleted = True
    if os.path.exists(docx_path):
        os.remove(docx_path)
        deleted = True
    
    if not deleted:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Generated cover letter not found",
        )
    
    return {"status": "deleted"}


@router.post("/{profile_id}/work-documents")
async def upload_work_documents(
    profile_id: str,
    work_experience_index: int,
    files: List[UploadFile] = File(...),
    db: AsyncSession = Depends(get_db),
):
    from app.services.document_parser import DocumentParser
    from sqlalchemy.orm.attributes import flag_modified
    import re
    
    profile = await get_profile_or_404(db, profile_id)
    work_exp = profile.work_experience or []
    validate_work_experience_index(work_exp, work_experience_index)

    # Validate file types
    allowed_types = [
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "text/markdown",
    ]

    # Create directory for work documents using user's name
    safe_name = re.sub(r'[^\w\s-]', '', profile.name).strip().replace(' ', '_')
    if not safe_name:
        safe_name = profile_id
    
    work_entry = work_exp[work_experience_index]
    company_name = work_entry.get("company_name", f"company_{work_experience_index}")
    safe_company = re.sub(r'[^\w\s-]', '', company_name).strip().replace(' ', '_')
    
    docs_dir = os.path.join(settings.storage_path, "work_documents", safe_name, safe_company)
    os.makedirs(docs_dir, exist_ok=True)

    uploaded_paths = []
    uploaded_contents = []
    
    for file in files:
        if file.content_type not in allowed_types:
            continue

        file_bytes = await file.read()
        
        # Keep original filename, handle duplicates
        original_filename = file.filename or "document.pdf"
        file_name = original_filename
        file_path = os.path.join(docs_dir, file_name)
        
        counter = 1
        name_part, ext_part = os.path.splitext(original_filename)
        while os.path.exists(file_path):
            file_name = f"{name_part}_{counter}{ext_part}"
            file_path = os.path.join(docs_dir, file_name)
            counter += 1

        # Save file to disk
        with open(file_path, "wb") as f:
            f.write(file_bytes)

        uploaded_paths.append(file_path)
        
        # Parse document content
        try:
            parsed_content, format_type = DocumentParser.parse_bytes(file_bytes, file_name)
            uploaded_contents.append({
                "filename": file_name,
                "path": file_path,
                "content": parsed_content,
                "format_type": format_type
            })
        except Exception as e:
            # If parsing fails, store error message but don't fail the upload
            uploaded_contents.append({
                "filename": file_name,
                "path": file_path,
                "content": f"[Failed to parse: {str(e)}]",
                "format_type": "error"
            })

    # Update work experience with document paths and contents
    updated_work_exp = []
    for i, entry in enumerate(work_exp):
        if i == work_experience_index:
            existing_paths = entry.get("document_paths", [])
            existing_contents = entry.get("document_contents", [])
            entry["document_paths"] = existing_paths + uploaded_paths
            entry["document_contents"] = existing_contents + uploaded_contents
        updated_work_exp.append(entry)
    
    profile.work_experience = updated_work_exp
    flag_modified(profile, "work_experience")

    await db.flush()
    await db.refresh(profile)
    
    work_entry = updated_work_exp[work_experience_index]

    return {
        "message": f"Uploaded and parsed {len(uploaded_paths)} documents",
        "document_paths": uploaded_paths,
        "document_contents": uploaded_contents,
        "work_experience": work_entry,
    }


@router.delete("/{profile_id}/work-documents/{work_experience_index}")
async def delete_work_document(
    profile_id: str,
    work_experience_index: int,
    document_path: str,
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy.orm.attributes import flag_modified
    
    profile = await get_profile_or_404(db, profile_id)
    work_exp = profile.work_experience or []
    validate_work_experience_index(work_exp, work_experience_index)

    # Create a deep copy of the work entry to trigger SQLAlchemy change detection
    work_entry = dict(work_exp[work_experience_index])
    doc_paths = list(work_entry.get("document_paths", []))
    doc_contents = list(work_entry.get("document_contents", []))
    
    if document_path in doc_paths:
        doc_paths.remove(document_path)
        # Also remove from document_contents
        doc_contents = [dc for dc in doc_contents if dc.get("path") != document_path]
        # Delete file from disk
        if os.path.exists(document_path):
            os.remove(document_path)
    
    work_entry["document_paths"] = doc_paths
    work_entry["document_contents"] = doc_contents
    
    # Create a NEW list to trigger SQLAlchemy's dirty detection
    updated_work_exp = []
    for i, entry in enumerate(work_exp):
        if i == work_experience_index:
            updated_work_exp.append(work_entry)
        else:
            updated_work_exp.append(dict(entry))  # Copy each entry
    
    # Assign new list
    profile.work_experience = updated_work_exp
    
    # Force the attribute to be marked as modified
    flag_modified(profile, "work_experience")

    await db.flush()
    await db.refresh(profile)

    return {
        "message": "Document deleted", 
        "document_paths": doc_paths,
        "work_experience": updated_work_exp[work_experience_index]
    }


@router.get("/{profile_id}/work-documents/{work_experience_index}/content")
async def get_document_content(
    profile_id: str,
    work_experience_index: int,
    document_path: str,
    db: AsyncSession = Depends(get_db),
):
    profile = await get_profile_or_404(db, profile_id)
    work_exp = profile.work_experience or []
    validate_work_experience_index(work_exp, work_experience_index)

    work_entry = work_exp[work_experience_index]
    doc_contents = work_entry.get("document_contents", [])
    
    # Find the document content by path
    for doc in doc_contents:
        if doc.get("path") == document_path:
            return {
                "filename": doc.get("filename"),
                "path": doc.get("path"),
                "content": doc.get("content"),
                "format_type": doc.get("format_type"),
            }
    
    # If not found in stored contents, try to parse from file
    if os.path.exists(document_path):
        from app.services.document_parser import DocumentParser
        try:
            content, format_type = DocumentParser.parse_file(document_path)
            filename = os.path.basename(document_path)
            return {
                "filename": filename,
                "path": document_path,
                "content": content,
                "format_type": format_type,
            }
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to parse document: {str(e)}",
            )
    
    raise HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail="Document not found",
    )


@router.get("/{profile_id}/resume/file")
async def get_resume_file(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    from fastapi.responses import FileResponse
    
    profile = await get_profile_or_404(db, profile_id)
    
    if not profile.resume_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No resume uploaded for this profile",
        )
    
    if not os.path.exists(profile.resume_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Resume file not found",
        )
    
    filename = os.path.basename(profile.resume_path)
    ext = os.path.splitext(filename)[1].lower()
    
    media_type_map = {
        ".pdf": "application/pdf",
        ".doc": "application/msword",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    media_type = media_type_map.get(ext, "application/octet-stream")
    
    return FileResponse(
        path=profile.resume_path,
        media_type=media_type,
        headers={"Content-Disposition": f"inline; filename=\"{filename}\""},
    )


@router.get("/{profile_id}/stats", response_model=ProfileStats)
async def get_profile_stats(
    profile_id: str,
    db: AsyncSession = Depends(get_db),
):
    await get_profile_or_404(db, profile_id)
    return await _get_profile_stats(db, profile_id)


async def _get_profile_stats(db: AsyncSession, profile_id: str) -> ProfileStats:
    query = (
        select(JobApplication.status, func.count(JobApplication.id))
        .where(JobApplication.profile_id == profile_id)
        .group_by(JobApplication.status)
    )
    result = await db.execute(query)
    status_counts = {row[0]: row[1] for row in result.all()}

    return ProfileStats(
        total_applications=sum(status_counts.values()),
        pending=status_counts.get(JobStatus.PENDING.value, 0) +
                status_counts.get(JobStatus.QUEUED.value, 0),
        in_progress=status_counts.get(JobStatus.IN_PROGRESS.value, 0),
        applied=status_counts.get(JobStatus.APPLIED.value, 0),
        failed=status_counts.get(JobStatus.FAILED.value, 0),
        awaiting_action=sum(status_counts.get(s, 0) for s in JobStatus.awaiting_statuses()),
    )
