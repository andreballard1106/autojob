import os
import re
from typing import Tuple
from pathlib import Path
from io import BytesIO


class DocumentParser:
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md'}

    @classmethod
    def parse_file(cls, file_path: str) -> Tuple[str, str]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = Path(file_path).suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        with open(file_path, 'rb') as f:
            content = f.read()

        return cls.parse_bytes(content, os.path.basename(file_path))

    @classmethod
    def parse_bytes(cls, content: bytes, filename: str) -> Tuple[str, str]:
        ext = Path(filename).suffix.lower()
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if ext == '.pdf':
            return cls._parse_pdf_content(content), 'pdf'
        elif ext in ('.docx', '.doc'):
            return cls._parse_docx_content(content), 'docx'
        elif ext == '.md':
            return cls._parse_text_content(content), 'markdown'
        else:
            return cls._parse_text_content(content), 'text'

    @classmethod
    def _parse_pdf_content(cls, content: bytes) -> str:
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(BytesIO(content))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    if page_num > 1:
                        text_parts.append(f"\n--- Page {page_num} ---\n")
                    text_parts.append(page_text)
            
            return cls._clean_text('\n'.join(text_parts))
            
        except ImportError:
            return "[PDF parsing requires PyPDF2. Install with: pip install PyPDF2]"
        except Exception as e:
            return f"[Error parsing PDF: {str(e)}]"

    @classmethod
    def _parse_docx_content(cls, content: bytes) -> str:
        try:
            from docx import Document
            import zipfile
            
            bytes_io = BytesIO(content)
            if not zipfile.is_zipfile(bytes_io):
                return "[Invalid DOCX file: not a valid ZIP archive]"
            
            bytes_io.seek(0)
            
            try:
                doc = Document(bytes_io)
            except Exception as doc_error:
                return cls._parse_docx_fallback(content, str(doc_error))
            
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    try:
                        style_name = para.style.name if para.style else ""
                        if style_name.startswith('Heading'):
                            level = style_name.replace('Heading ', '')
                            try:
                                level_num = int(level)
                                prefix = '#' * min(level_num, 6) + ' '
                            except ValueError:
                                prefix = '## '
                            text_parts.append(f"\n{prefix}{para.text}\n")
                        else:
                            text_parts.append(para.text)
                    except Exception:
                        text_parts.append(para.text)
            
            try:
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            table_text.append(row_text)
                    if table_text:
                        text_parts.append('\n' + '\n'.join(table_text) + '\n')
            except Exception:
                pass
            
            return cls._clean_text('\n'.join(text_parts))
            
        except ImportError:
            return "[DOCX parsing requires python-docx. Install with: pip install python-docx]"
        except Exception as e:
            return f"[Error parsing DOCX: {str(e)}]"

    @classmethod
    def _parse_docx_fallback(cls, content: bytes, original_error: str) -> str:
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            bytes_io = BytesIO(content)
            
            with zipfile.ZipFile(bytes_io, 'r') as zf:
                if 'word/document.xml' not in zf.namelist():
                    return f"[Invalid DOCX structure: {original_error}]"
                
                xml_content = zf.read('word/document.xml')
                root = ET.fromstring(xml_content)
                
                text_parts = []
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                for para in root.iter(f'{ns}p'):
                    para_text = ''
                    for text_elem in para.iter(f'{ns}t'):
                        if text_elem.text:
                            para_text += text_elem.text
                    if para_text.strip():
                        text_parts.append(para_text)
                
                if text_parts:
                    return cls._clean_text('\n'.join(text_parts))
                else:
                    return f"[No text content found in DOCX: {original_error}]"
                    
        except Exception as e:
            return f"[Fallback DOCX parsing failed: {str(e)}. Original error: {original_error}]"

    @classmethod
    def _parse_text_content(cls, content: bytes) -> str:
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        try:
            import chardet
            detected = chardet.detect(content)
            return content.decode(detected['encoding'] or 'utf-8')
        except ImportError:
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            return f"[Error reading text: {str(e)}]"

    @classmethod
    def _clean_text(cls, text: str) -> str:
        if not text:
            return ""
        
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        lines = [line.rstrip() for line in text.split('\n')]
        return '\n'.join(lines).strip()

    @classmethod
    def get_content_preview(cls, content: str, max_length: int = 500) -> str:
        if len(content) <= max_length:
            return content
        return content[:max_length] + "..."
